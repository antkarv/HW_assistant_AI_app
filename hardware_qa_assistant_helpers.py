# hardware_qa_assistant_helpers.py
from __future__ import annotations

import json
import os, io, re, glob, tempfile, urllib.parse, time
from typing import List, Dict, Any, Tuple, Optional

# ---- Vector DB (persistent, local) ----
import chromadb
from chromadb.utils import embedding_functions

from chromadb.utils import embedding_functions

CHROMA_DIR = os.getenv("CHROMA_DIR", ".chroma_hwqa_workspace")
os.makedirs(CHROMA_DIR, exist_ok=True)
chroma = chromadb.PersistentClient(path=CHROMA_DIR)

# --- Demo toggle (keep HF, just choose smaller model when DEMO_MODE is on) ---
_DEMO = os.getenv("DEMO_MODE", "0").lower() in ("1", "true", "yes")
_EMB_DEFAULT = "sentence-transformers/all-mpnet-base-v2"
_EMB_DEMO    = "sentence-transformers/all-MiniLM-L6-v2"  # small, fast, great for demos
EMB_MODEL = os.getenv("EMB_MODEL") or (_EMB_DEMO if _DEMO else _EMB_DEFAULT)

_embed_fn = None
def get_embed_fn():
    """Create the embedding function on first use (saves memory + port binds fast)."""
    global _embed_fn
    if _embed_fn is None:
        # keep torch lean in small containers
        os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
        os.environ.setdefault("OMP_NUM_THREADS", "1")
        os.environ.setdefault("MKL_NUM_THREADS", "1")
        _embed_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=EMB_MODEL
        )
        print(f"[emb] Using {EMB_MODEL} (demo={_DEMO})")
    return _embed_fn


# ---- OCR + parsing stack (isolated here) ----
import fitz                        # PyMuPDF
import docx                        # python-docx
from pdf2image import convert_from_bytes
from PIL import Image
import cv2, numpy as np

# -------------------------
# Small workspace utilities
# -------------------------
def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (s or "").strip().lower()).strip("-") or "default"

def user_prefix(email: str) -> str:
    return f"u_{_slug(email)}__ws_"

def collection_name(email: Optional[str], workspace: str) -> str:
    return f"{user_prefix(email)}{_slug(workspace)}" if email else f"ws_{_slug(workspace)}"

def list_projects_for_user(email: str) -> List[str]:
    pref = user_prefix(email)
    try:
        cols = chroma.list_collections()
    except Exception:
        return []
    out = []
    for c in cols:
        if c.name.startswith(pref):
            out.append(c.name[len(pref):])
    return sorted(set(out))

def delete_project(email: str, name: str) -> bool:
    try:
        chroma.delete_collection(f"{user_prefix(email)}{_slug(name)}")
        return True
    except Exception:
        return False

# -------------------------
# Chroma helpers
# -------------------------
def get_collection(workspace: str, user: Optional[str] = None):
    name = collection_name(user, workspace)
    try:
        # bind the embedding function even when getting an existing collection
        return chroma.get_collection(name, embedding_function=get_embed_fn())
    except Exception:
        return chroma.create_collection(name=name, embedding_function=get_embed_fn())


def add_docs(col, docs: List[Tuple[str, str, Dict[str, Any]]]):
    if not docs:
        return
    ids = [d[0] for d in docs]
    texts = [d[1] for d in docs]
    metas = [d[2] for d in docs]
    col.add(ids=ids, documents=texts, metadatas=metas)

# def retrieve(col, query: str, k: int = 5) -> List[str]:
#     if not query.strip():
#         return []
#     res = col.query(query_texts=[query], n_results=k)
#     return res.get("documents", [[]])[0]

def retrieve(col, query: str, k: int = 5) -> List[str]:
    docs, _, _, _ = retrieve_info(col, query, k)
    return docs

def retrieve_info(col, query: str, k: int = 5):
    """Return (docs, metadatas, distances, ids) for debugging/inspection."""
    if not query.strip():
        return [], [], [], []
    res = col.query(
        query_texts=[query],
        n_results=k,
        include=["documents", "metadatas", "distances"],
    )
    docs = res.get("documents", [[]])[0]
    metas = res.get("metadatas", [[]])[0]
    dists = res.get("distances", [[]])[0]
    ids   = res.get("ids", [[]])[0]
    return docs, metas, dists, ids


# ---------- Retrieval quality utilities (NEW) ----------

from typing import Iterable
import math

def rrf_fuse(rank_lists: List[List[Tuple[str, Dict[str,Any], float, str]]], k: int = 60, K: int = 60):
    """
    rank_lists: list over sub-queries; each is [(doc, meta, dist, id), ...] sorted by DB rank.
    Return fused unique list [(doc, meta, fused_score, id)] sorted desc by fused_score.
    K is the RRF 'rank constant' (larger → flatter).
    """
    scores = {}
    seen_meta = {}
    for lst in rank_lists:
        for r, (doc, meta, dist, _id) in enumerate(lst, start=1):
            key = _id or (doc[:60] + "|" + (meta or {}).get("filename",""))
            scores[key] = scores.get(key, 0.0) + 1.0 / (K + r)
            if key not in seen_meta:
                seen_meta[key] = (doc, meta, (1.0 - float(dist)) if dist is not None else None, _id)
    fused = [(seen_meta[k][0], seen_meta[k][1], scores[k], seen_meta[k][3]) for k in scores]
    fused.sort(key=lambda x: x[2], reverse=True)
    return fused[:k]

def _embed_texts(texts: List[str]) -> List[List[float]]:
    # reuse the same embedding fn as Chroma uses
    ef = get_embed_fn()
    return ef(texts)


from functools import lru_cache

@lru_cache(maxsize=1)
def _load_cross_encoder(model_name: str = "cross-encoder/ms-marco-MiniLM-L-6-v2"):
    from sentence_transformers import CrossEncoder
    return CrossEncoder(model_name)

def rerank_cross_encoder(query: str, cand_docs: List[Tuple[str,Dict[str,Any],float,str]],
                         model_name: str = "cross-encoder/ms-marco-MiniLM-L-6-v2") -> List[Tuple[str,Dict[str,Any],float,str]]:
    """
    Return same tuples [(doc, meta, score, id)] but with cross-encoder similarity score, sorted desc.
    Falls back to dot(query_emb, doc_emb) if CrossEncoder not available.
    """
    try:
        ce = _load_cross_encoder(model_name)
        pairs = [(query, d[0]) for d in cand_docs]
        scores = ce.predict(pairs).tolist()
        out = [(d[0], d[1], float(s), d[3]) for d, s in zip(cand_docs, scores)]
        out.sort(key=lambda x: x[2], reverse=True)
        return out
    except Exception:
        # fallback: cosine with same embedding model used for the DB
        vecs = _embed_texts([query] + [d[0] for d in cand_docs])
        qv, dvs = vecs[0], vecs[1:]
        def _cos(a,b):
            import numpy as _np
            a = _np.array(a); b = _np.array(b)
            denom = (float((a*a).sum())**0.5) * (float((b*b).sum())**0.5)
            return float((a*b).sum()) / denom if denom else 0.0
        out = [(d[0], d[1], _cos(qv, dv), d[3]) for d, dv in zip(cand_docs, dvs)]
        out.sort(key=lambda x: x[2], reverse=True)
        return out

def mmr_select(query: str, cand_docs: List[Tuple[str,Dict[str,Any],float,str]],
               keep: int = 8, diversity_lambda: float = 0.7) -> List[Tuple[str,Dict[str,Any],float,str]]:
    """
    Maximal Marginal Relevance on the (already re-ranked) candidate list to keep diverse top-N.
    """
    vecs = _embed_texts([query] + [d[0] for d in cand_docs])
    qv, dvs = vecs[0], vecs[1:]
    import numpy as _np
    def _cos(a,b):
        a = _np.array(a); b = _np.array(b)
        denom = (float((a*a).sum())**0.5) * (float((b*b).sum())**0.5)
        return float((a*b).sum()) / denom if denom else 0.0

    selected, rest = [], list(range(len(cand_docs)))
    while rest and len(selected) < keep:
        if not selected:
            # start from the best already (list is sorted)
            best = rest.pop(0)
            selected.append(best)
            continue
        # pick argmax of λ*sim(query,di) - (1-λ)*max_j sim(di, dj_selected)
        best_idx, best_score = None, -1e9
        for i in rest:
            rel = _cos(qv, dvs[i])
            red = max(_cos(dvs[i], dvs[j]) for j in selected) if selected else 0.0
            score = diversity_lambda * rel - (1.0 - diversity_lambda) * red
            if score > best_score:
                best_idx, best_score = i, score
        rest.remove(best_idx)
        selected.append(best_idx)
    return [cand_docs[i] for i in selected]

def _normalize_where(where: Dict[str, Any]) -> Dict[str, Any]:
    """
    Chroma (newer) expects a single top-level operator: $and / $or / $not.
    Convert a simple dict like {"filename":"x","doc_type":"pdf","page":2}
    into {"$and":[{"filename":{"$eq":"x"}},{"doc_type":{"$eq":"pdf"}},{"page":{"$eq":2}}]}.
    If already operator-structured, return as-is.
    """
    if not where:
        return {}
    # if already has a top-level operator, keep it
    if any(str(k).startswith("$") for k in where.keys()):
        return where
    clauses = []
    for k, v in where.items():
        if v is None:
            continue
        if isinstance(v, dict) and any(str(op).startswith("$") for op in v.keys()):
            clauses.append({k: v})
        else:
            clauses.append({k: {"$eq": v}})
    return {"$and": clauses} if clauses else {}


def get_by_where(col, where: Dict[str, Any]) -> Tuple[List[str], List[Dict[str,Any]]]:
    where = _normalize_where(where)   # <-- add this line
    res = col.get(where=where, include=["documents","metadatas"])
    docs = res.get("documents") or []
    metas = res.get("metadatas") or []
    return docs, metas

def promote_parents(col, picked: List[Tuple[str,Dict[str,Any],float,str]], max_siblings_per_parent: int = 1) -> List[Tuple[str,Dict[str,Any],float,str]]:
    """
    For each picked child chunk, fetch siblings from the same parent (PDF page or DOCX section).
    """
    out: List[Tuple[str,Dict[str,Any],float,str]] = []
    seen_ids = set()
    for (doc, meta, score, _id) in picked:
        out.append((doc, meta, score, _id))
        parent_where = None
        m = meta or {}
        fn = m.get("filename") or m.get("source") or m.get("name")
        if not fn:
            continue
        if m.get("doc_type") == "pdf" and m.get("page") is not None:
            parent_where = {"filename": fn, "doc_type": "pdf", "page": m.get("page")}
        elif m.get("doc_type") == "docx" and m.get("section_path"):
            parent_where = {"filename": fn, "doc_type": "docx", "section_path": m.get("section_path")}
        if not parent_where:
            continue
        sdocs, smetas = get_by_where(col, parent_where)
        added = 0
        for sd, sm in zip(sdocs, smetas):
            sid = (sd[:60] + "|" + (sm or {}).get("filename","") + "|" + str((sm or {}).get("chunk")))
            if sid in seen_ids:
                continue
            if added >= max_siblings_per_parent:
                break
            # avoid re-adding the very same child (best effort without global ids)
            if sd.strip() == (doc or "").strip():
                continue
            out.append((sd, sm, score, sid))
            seen_ids.add(sid)
            added += 1
    return out



# -------------------------
# OCR backends + parsers
# -------------------------
_PADDLE_OCR = None
def _get_paddle_ocr():
    global _PADDLE_OCR
    if _PADDLE_OCR is not None:
        return _PADDLE_OCR
    try:
        from paddleocr import PaddleOCR
        _PADDLE_OCR = PaddleOCR(use_angle_cls=True, lang="en")
    except Exception as e:
        _PADDLE_OCR = f"[paddle init error] {e}"
    return _PADDLE_OCR

def ocr_image_paddle(img: Image.Image) -> str:
    ocr = _get_paddle_ocr()
    if isinstance(ocr, str):
        return ocr
    arr = np.array(img.convert("RGB"))[:, :, ::-1]  # BGR
    result = ocr.ocr(arr)
    lines: List[str] = []
    for page in result or []:
        for item in page or []:
            try:
                _, (txt, _conf) = item
                if txt:
                    lines.append(txt)
            except Exception:
                pass
    return "\n".join(lines).strip()

def ocr_image_tesseract(img: Image.Image) -> str:
    try:
        import pytesseract
        gray = cv2.cvtColor(np.array(img.convert("RGB")), cv2.COLOR_RGB2GRAY)
        thr = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        return pytesseract.image_to_string(Image.fromarray(thr)).strip()
    except Exception as e:
        return f"[tesseract error] {e}"

def ocr_image(img: Image.Image, backend: str) -> str:
    b = (backend or "auto").lower()
    if b == "paddle":
        out = ocr_image_paddle(img)
        return out if out and not out.startswith("[paddle init error]") else ocr_image_tesseract(img)
    if b == "tesseract":
        return ocr_image_tesseract(img)
    out = ocr_image_paddle(img)
    return out if out and not out.startswith("[paddle init error]") else ocr_image_tesseract(img)

def parse_pdf(file_bytes: bytes, *, ocr_backend: str) -> str:
    text_parts: List[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for p in doc:
            text_parts.append(p.get_text("text"))
    text = "\n".join(text_parts).strip()
    if text:
        return text
    pages = convert_from_bytes(file_bytes)
    lines = [ocr_image(im, backend=ocr_backend) for im in pages]
    return "\n".join(lines)

def parse_docx(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    d = docx.Document(path)
    os.unlink(path)
    return "\n".join(p.text for p in d.paragraphs)

def parse_image(file_bytes: bytes, *, ocr_backend: str) -> str:
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    return ocr_image(img, backend=ocr_backend)

# -------------------------
# Hybrid chunking helpers (PDF pages + DOCX sections)  ⬇️  ADD THIS
# -------------------------
def chunk_text(s: str, size: int = 1200, overlap: int = 180) -> List[str]:
    """Slide a window over s with the given size/overlap, returning chunks."""
    s = (s or "").strip()
    out: List[str] = []
    if not s:
        return out
    i = 0
    n = len(s)
    while i < n:
        j = min(n, i + size)
        out.append(s[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)
    return out

def parse_pdf_pages(file_bytes: bytes, *, ocr_backend: str) -> List[str]:
    """Return text per page (OCRing only pages that have no extractable text)."""
    texts: List[str] = []
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            page_count = len(doc)
            for idx in range(page_count):
                t = (doc[idx].get_text("text") or "").strip()
                if t:
                    texts.append(t)
                else:
                    # OCR just this page to avoid rendering whole PDF into RAM
                    try:
                        imgs = convert_from_bytes(
                            file_bytes, first_page=idx + 1, last_page=idx + 1
                        )
                        t = ocr_image(imgs[0], backend=ocr_backend)
                    except Exception as e:
                        t = f"[ocr error p{idx+1}] {e}"
                    texts.append(t or "")
    except Exception as e:
        # fallback: keep previous parse_pdf behavior (may OCR all pages)
        texts = (parse_pdf(file_bytes, ocr_backend=ocr_backend) or "").split("\n\f\n")
    return texts

def parse_docx_sections(file_bytes: bytes) -> List[Tuple[str, str]]:
    """Return (section_path, text) pairs using Heading 1/2/3 as boundaries.
       If no headings exist, return a single ('Body', full_text).
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    d = docx.Document(path)
    os.unlink(path)

    def _lvl(name: Optional[str]) -> Optional[int]:
        if not name:
            return None
        m = re.match(r"heading\s*(\d+)", name.strip().lower())
        return int(m.group(1)) if m else None

    sections: List[Tuple[str, str]] = []
    stack: List[str] = []     # e.g., ["Intro", "Clocking"]
    buf: List[str] = []

    def _flush():
        if buf:
            sec = " > ".join(stack) if stack else "Body"
            sections.append((sec, "\n".join(buf).strip()))
            buf.clear()

    for p in d.paragraphs:
        txt = (p.text or "").strip()
        lvl = _lvl(getattr(getattr(p, "style", None), "name", None))
        if lvl and 1 <= lvl <= 3:
            _flush()
            # adjust heading path depth
            if lvl == 1:
                stack = [txt] if txt else []
            elif lvl == 2:
                stack = (stack[:1] + [txt]) if txt else stack[:1]
            else:  # lvl == 3
                stack = (stack[:2] + [txt]) if txt else stack[:2]
        else:
            if txt:
                buf.append(txt)
    _flush()

    if not sections:
        whole = "\n".join(p.text for p in d.paragraphs).strip()
        return [("Body", whole)]
    return sections



# -------------------------
# Manage-PDF helpers
# -------------------------
def list_workspace_files_by_filename(ws: str, email: Optional[str]) -> List[str]:
    if not ws or not email:
        return []
    col = get_collection(ws, email)
    res = col.get(include=["metadatas"])
    metadatas = res.get("metadatas") or []
    names: List[str] = []
    for m in metadatas:
        m = m or {}
        fn = m.get("source") or m.get("filename") or m.get("name")
        if fn:
            names.append(os.path.basename(str(fn)))
    return sorted(set(names), key=str.lower)

def delete_files_by_filename(ws: str, email: Optional[str], filenames: List[str]) -> str:
    if not ws:
        return "Select a project first."
    if not email:
        return "You must be logged in."
    if not filenames:
        return "Nothing selected."
    targets = {os.path.basename(str(n)) for n in filenames}
    col = get_collection(ws, email)
    res = col.get(include=["metadatas"])
    metadatas = res.get("metadatas") or []
    ids = res.get("ids") or []
    to_delete_ids: List[str] = []
    for _id, m in zip(ids, metadatas):
        m = m or {}
        fn = m.get("source") or m.get("filename") or m.get("name")
        if fn and os.path.basename(str(fn)) in targets:
            to_delete_ids.append(_id)
    if not to_delete_ids:
        return "No matching chunks found for the selected filename(s)."
    col.delete(ids=to_delete_ids)
    return f"Deleted {len(to_delete_ids)} chunk(s) from project '{ws}' across {len(targets)} file name(s)."

# -------------------------
# Reports helpers
# -------------------------
def list_summary_reports() -> List[str]:
    roots = [os.getcwd(), tempfile.gettempdir()]
    out: List[str] = []
    for root in roots:
        out.extend(os.path.abspath(p) for p in glob.glob(os.path.join(root, "hwqa_summary_*.md")))
    out.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return out

def read_report_text(path: str) -> str:
    if not path:
        return "_No report selected._"
    if not os.path.isfile(path):
        return f"_Report not found:_ `{path}`"
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"_Could not read report (`{path}`):_ {e}"

# -------------------------
# Trace helpers
# -------------------------
def trace_push(state: Dict[str, Any], node: str, t0: float, notes: str = "") -> None:
    try:
        ms = int((time.perf_counter() - t0) * 1000)
    except Exception:
        ms = 0
    trace = state.get("trace") or []
    trace.append({"node": node, "ms": ms, "notes": notes})
    state["trace"] = trace

def _trace_to_markdown(state: Dict[str, Any]) -> str:
    q = (state.get("user_input") or "").strip()
    full = state.get("trace") or []

    # keep only entries from the last 'session' marker onward
    last_idx = -1
    for i, e in enumerate(full):
        if (e or {}).get("node") == "session":
            last_idx = i
    trace = full[last_idx:] if last_idx != -1 else full

    head = [
        "### Last request",
        f"**Q:** {q}" if q else "**Q:** (no question text)",
        f"Workspace: {state.get('workspace','')}   OCR: {state.get('ocr_backend','')}   Web search: {state.get('do_web_search')}   Schematic mode: {state.get('treat_images_as_schematics')}",
        "",
    ]
    body = []
    for i, e in enumerate(trace or [], 1):
        body.append(f"{i}. **{e.get('node','?')}** — {e.get('notes','')} ({e.get('ms',0)} ms)")
    cites = []
        # Show citations only if the web node actually ran and returned sources>0
    # (prevents lingering/irrelevant links when web search is disabled)
    import re as _re

    # work on the sliced "trace" we already computed above
    web_sources = 0
    for e in trace or []:
        if (e or {}).get("node") == "web":
            notes = (e or {}).get("notes", "")
            m = _re.search(r"sources\s*=\s*(\d+)", notes)
            if m:
                try:
                    web_sources = int(m.group(1))
                except Exception:
                    web_sources = 0

    if web_sources > 0:
        cites = [f"- {u}" for u in (state.get("citations") or [])]
        foot = ["", "**Citations (web):**", *cites]
    else:
        foot = []


    # ---- LLM Inspector (NEW) ----
    events = state.get("llm_events") or []
    inspector: List[str] = []
    if events:
        inspector.extend(["", "### LLM Inspector", "",
                          "| node | model | ms | in | out | total |",
                          "|---|---|---:|---:|---:|---:|"])
        for e in events:
            inspector.append(
                f"| {e.get('node','')} | {e.get('model','')} | {e.get('latency_ms',0)} | "
                f"{e.get('prompt_tokens',0)} | {e.get('completion_tokens',0)} | {e.get('total_tokens',0)} |"
            )
        inspector.append("")
        inspector.append("<details><summary>Show input/output previews</summary>")
        for e in events:
            inspector.append(
                f"<br/><b>{e.get('node','')}</b> — <i>{e.get('model','')}</i><br/>"
                f"<b>Input:</b> {e.get('prompt_preview','')}<br/>"
                f"<b>Output:</b> {e.get('output_preview','')}"
            )
        inspector.append("</details>")
        # ---- Retrieval Inspector (NEW) ----
    r_events = state.get("retrieval_events") or []
    retr_tbl: List[str] = []
    if r_events:
        retr_tbl.extend(["", "### Retrieval Inspector", "",
                         "| rank | score | file | type | page/section | preview |",
                         "|---:|---:|---|---|---|---|"])
        for e in r_events:
            loc = ""
            if e.get("page") is not None:
                loc = f"p{e.get('page')}"
            elif e.get("section"):
                loc = e.get("section") or ""
            # escape pipes in preview
            prev = (e.get("preview","") or "").replace("|", "\\|")
            score = "" if e.get("score") is None else f"{e.get('score'):.3f}"
            retr_tbl.append(f"| {e.get('rank','')} | {score} | {e.get('filename','')} | "
                            f"{e.get('doc_type','')} | {loc} | {prev} |")

        # ---- Executed Nodes Summary + Execution Narrative (NEW) ----
    trace_list = state.get("trace") or []
    ran = [e for e in trace_list if e.get("node")]
    ran_names = [e.get("node") for e in ran]

    # helpers to pick info from state
    def _llm_line(node_name: str):
        for e in (state.get("llm_events") or []):
            if e.get("node") == node_name:
                m = e.get("model") or ""
                pi = e.get("prompt_tokens") or 0
                co = e.get("completion_tokens") or 0
                ms = e.get("latency_ms") or 0
                return f"model={m}  in={pi}  out={co}  {int(ms)}ms"
        return ""

    def _retrieve_line():
        evs = state.get("retrieval_events") or []
        if not evs:
            return "hits=0"
        top = ", ".join([f"{(e.get('filename') or '')} {('p'+str(e['page'])) if e.get('page') else (e.get('section') or '')}".strip()
                         for e in evs[:3]])
        return f"hits={len(evs)} ({top})"

    summary_lines: List[str] = []
    for idx, e in enumerate(ran, 1):
        n = e.get("node", "?")
        notes = e.get("notes", "")
        if n == "router":
            route = state.get("route") or "?"
            files = len(state.get("parsed_docs") or [])
            schem = bool(state.get("has_schematic"))
            summary_lines.append(f"{idx}. **router** — route={route} (files={files}, schematic={schem})")
        elif n == "ingest":
            summary_lines.append(f"{idx}. **ingest** — {notes or 'done'}")
        elif n == "retrieve":
            # Legacy path (older runs) — keep this
            k = len(state.get("retrieved_chunks") or [])
            summary_lines.append(f"{idx}. **retrieve** — returned {k} chunks")
        elif n in ("multiquery","retrieve_pool","rrf_fuse","rerank","mmr","parent_promote"):
            if n == "multiquery":
                summary_lines.append(f"{idx}. **multiquery** — generated variants")
            elif n == "retrieve_pool":
                summary_lines.append(f"{idx}. **retrieve_pool** — pooled top-k per variant")
            elif n == "rrf_fuse":
                summary_lines.append(f"{idx}. **rrf_fuse** — fused & deduped")
            elif n == "rerank":
                summary_lines.append(f"{idx}. **rerank** — cross-encoder re-ordered")
            elif n == "mmr":
                summary_lines.append(f"{idx}. **mmr** — selected diverse top-N")
            elif n == "parent_promote":
                # New pipeline — final point where retrieved_chunks is set
                k = len(state.get("retrieved_chunks") or [])
                summary_lines.append(f"{idx}. **parent_promote** — finalized {k} chunks for context")

        elif n == "web":
            ws = state.get("web_snippets") or []
            summary_lines.append(f"{idx}. **web** — sources={len(ws)}")
        elif n in ("answer", "schematic_answer"):
            summary_lines.append(f"{idx}. **{n}** — {_llm_line(n)}")
        elif n == "inventory_docs":
            summary_lines.append(f"{idx}. **inventory_docs** — {notes or 'listed PDFs/DOCX'}")
        elif n == "inventory_schematics":
            summary_lines.append(f"{idx}. **inventory_schematics** — {notes or 'listed schematics'}")
        elif n == "verify":
            summary_lines.append(f"{idx}. **verify** — {notes or 'ok'}")
        else:
            summary_lines.append(f"{idx}. **{n}** — {notes}")

    # Build a one-paragraph narrative
    narrative_bits: List[str] = []
    if "router" in ran_names:
        narrative_bits.append(f"The router selected the **{state.get('route') or '?'}** path.")
    if "ingest" in ran_names:
        narrative_bits.append("Ingest parsed and indexed the new files.")
    if "retrieve" in ran_names:
        evs = state.get("retrieval_events") or []
        narrative_bits.append(f"Retrieval returned **{len(evs)}** chunk(s).")
    if "web" in ran_names:
        ws = state.get("web_snippets") or []
        if ws:
            narrative_bits.append(f"Web search added **{len(ws)}** source(s).")
        else:
            narrative_bits.append("No web sources were needed.")
    if "answer" in ran_names or "schematic_answer" in ran_names:
        # use whichever LLM event we have
        llms = state.get("llm_events") or []
        if llms:
            m = llms[-1].get("model", "")
            pi = llms[-1].get("prompt_tokens") or 0
            co = llms[-1].get("completion_tokens") or 0
            narrative_bits.append(f"The model **{m}** generated the answer ({pi} prompt → {co} completion tokens).")
    if "verify" in ran_names:
        narrative_bits.append("Verification passed.")

    summary_block: List[str] = []
    if summary_lines:
        summary_block.extend(["", "### What happened in this run", ""])
        summary_block.extend([*summary_lines, ""])
    if narrative_bits:
        summary_block.extend(["**Execution narrative:** " + " ".join(narrative_bits), ""])

    # ---- Raw JSON collapsed (NEW) ----
    raw = {"trace": trace_list}
    raw_json = json.dumps(raw, indent=2, ensure_ascii=False)
    raw_block = [
        "",
        "<details><summary><b>Show raw trace JSON</b></summary>",
        "",
        "```json",
        raw_json,
        "```",
        "",
        "</details>",
    ]

    return "\n".join(head + body + inspector + retr_tbl + summary_block + raw_block + foot)

def _dag_image_from_trace(trace: list, llm_events: list = None):
    import networkx as nx
    import matplotlib.pyplot as plt
    from PIL import Image as PILImage
    import io as _io

    G = nx.DiGraph()

    nodes = [
        "session", "router", "ingest", 
        "multiquery","retrieve_pool","rrf_fuse","rerank","mmr","parent_promote",
        "web", "answer", "verify", "clarify", "report", "END",
        "parts_expand", "schematic_answer",
        "inventory_docs", "inventory_schematics"
    ]
    G.add_nodes_from(nodes)

    G.add_edges_from([
        ("session","router"), ("session","report"),
        ("router","ingest"), ("router","multiquery"),
        ("ingest","multiquery"),
        ("multiquery","retrieve_pool"),
        ("retrieve_pool","rrf_fuse"),
        ("rrf_fuse","rerank"),
        ("rerank","mmr"),
        ("mmr","parent_promote"),

        ("parent_promote","web"),
        ("parent_promote","answer"),
        ("web","answer"),
        ("answer","verify"),
        ("parent_promote","parts_expand"),
        ("parts_expand","schematic_answer"),
        ("schematic_answer","verify"),
        ("verify","clarify"),
        ("verify","END"),
        ("parent_promote","inventory_docs"),
        ("parent_promote","inventory_schematics"),
        ("inventory_docs","verify"),
        ("inventory_schematics","verify"),
    ])

    pos = {
        "session": (0.0, 0.0),
        "router":  (1.2, 0.0),
        "ingest":  (2.4, 0.6),
        "multiquery":     (3.6, 0.6),
        "retrieve_pool":  (4.8, 0.6),
        "rrf_fuse":       (6.0, 0.6),
        "rerank":         (7.2, 0.6),
        "mmr":            (8.4, 0.6),
        "parent_promote": (9.6, 0.6),

        "web":    (10.8, 1.0),
        "answer": (12.0, 0.6),
        "verify": (13.2, 0.6),
        "clarify":(14.4, 1.0),
        "report": (2.4, -1.0),
        "END":    (14.4, 0.2),

        "parts_expand":        (10.8, 0.0),
        "schematic_answer":    (12.0, 0.2),
        "inventory_docs":      (10.8, 0.4),
        "inventory_schematics":(10.8, -0.4),
    }

    # Map latest LLM event per node
    ev_map = {}
    for e in (llm_events or []):
        ev_map[e.get("node")] = e

    # Highlight executed path
    path_nodes = [e.get("node") for e in (trace or []) if e.get("node") in G.nodes]
    path_edges = [(a, b) for a, b in zip(path_nodes, path_nodes[1:]) if G.has_edge(a, b)]

    node_colors = ["#90caf9" if n in path_nodes else "#e6e6e6" for n in G.nodes]
    node_sizes  = [1800 if n in path_nodes else 1200 for n in G.nodes]
    edge_colors = ["#1976d2" if e in path_edges else "#bdbdbd" for e in G.edges]
    widths      = [3.0 if e in path_edges else 1.5 for e in G.edges]

    # Labels with subtitle if LLM used
    labels = {}
    for n in nodes:
        if n in ev_map:
            e = ev_map[n]
            labels[n] = f"{n}\n{e.get('model','')} · {e.get('prompt_tokens',0)}→{e.get('completion_tokens',0)} tok"
        else:
            labels[n] = n

    fig = plt.figure(figsize=(11, 3.6))
    nx.draw(G, pos, with_labels=False, arrows=True,
            node_color=node_colors, node_size=node_sizes,
            edge_color=edge_colors, width=widths)
    nx.draw_networkx_labels(G, pos, labels=labels, font_size=9)
    plt.axis("off")
    buf = _io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150); plt.close(fig)
    buf.seek(0)
    return PILImage.open(buf)

def render_trace_artifacts(state: Dict[str, Any]):
    md = _trace_to_markdown(state)
    json_str = __import__("json").dumps(state.get("trace") or [], indent=2)

    # last-turn slice for the DAG too
    full = state.get("trace") or []
    last_idx = -1
    for i, e in enumerate(full):
        if (e or {}).get("node") == "session":
            last_idx = i
    last_trace = full[last_idx:] if last_idx != -1 else full

    try:
        img = _dag_image_from_trace(last_trace, state.get("llm_events") or [])
    except Exception:
        img = None
    return md, json_str, img

